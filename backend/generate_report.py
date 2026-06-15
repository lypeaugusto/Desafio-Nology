from docx import Document
from cashback import compute_from_price

def formato_real(v):
    return f"R$ {v:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

def main():
    doc = Document()
    doc.add_heading('Respostas - Cashback', level=1)

    # Caso 2: VIP, preço 600, cupom 20%
    r2 = compute_from_price(600, 20, True)
    doc.add_heading('2) VIP - R$ 600 com 20% off', level=2)
    doc.add_paragraph(f"Preço original: {formato_real(r2['preco_original'])}")
    doc.add_paragraph(f"Desconto: {r2['desconto_percent']}% -> Valor final: {formato_real(r2['valor_final'])}")
    doc.add_paragraph(f"Cashback base (5%): {formato_real(r2['cashback_base'])}")
    doc.add_paragraph(f"Bônus VIP (10% sobre base): {formato_real(r2['cashback_bonus'])}")
    doc.add_paragraph(f"Promoção (valor final > 500): cashback dobrado -> Cashback final: {formato_real(r2['cashback_final'])}")

    # Caso 3: Normal, preço 600, cupom 10%
    r3 = compute_from_price(600, 10, False)
    doc.add_heading('3) Normal - R$ 600 com 10% off', level=2)
    doc.add_paragraph(f"Preço original: {formato_real(r3['preco_original'])}")
    doc.add_paragraph(f"Desconto: {r3['desconto_percent']}% -> Valor final: {formato_real(r3['valor_final'])}")
    doc.add_paragraph(f"Cashback base (5%): {formato_real(r3['cashback_base'])}")
    doc.add_paragraph(f"Sem bônus VIP. Promoção (valor final > 500): cashback dobrado -> Cashback final: {formato_real(r3['cashback_final'])}")

    # Caso 4: Reclamação - VIP 600 com 15% off, recebeu R$56
    r4 = compute_from_price(600, 15, True)
    doc.add_heading('4) Reclamação - VIP - R$ 600 com 15% off', level=2)
    doc.add_paragraph(f"Preço original: {formato_real(r4['preco_original'])}")
    doc.add_paragraph(f"Desconto: {r4['desconto_percent']}% -> Valor final: {formato_real(r4['valor_final'])}")
    doc.add_paragraph(f"Cashback base (5%): {formato_real(r4['cashback_base'])}")
    doc.add_paragraph(f"Bônus VIP (10% sobre base): {formato_real(r4['cashback_bonus'])}")
    doc.add_paragraph(f"Promoção (valor final > 500): cashback dobrado -> Cashback final esperado: {formato_real(r4['cashback_final'])}")
    doc.add_paragraph('')
    doc.add_paragraph('O cliente afirmou ter recebido R$ 56. Abaixo justificamos o valor esperado:')
    doc.add_paragraph(f"Esperado: {formato_real(r4['cashback_final'])}. Recebido (reportado): R$ 56,00.")
    if abs(r4['cashback_final'] - 56.0) > 0.01:
        doc.add_paragraph('Conclusão: O valor recebido está incorreto de acordo com as regras atuais. Por favor, abra investigação para verificar possível problema no processamento ou arredondamento.')
    else:
        doc.add_paragraph('Conclusão: O valor recebido está correto de acordo com as regras atuais.')

    doc.save('respostas_cashback.docx')
    print('Documento gerado: respostas_cashback.docx')

if __name__ == '__main__':
    main()
